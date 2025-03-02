# to handle file operations such as reading, writing, and listing files.
# Supports TXT, PDF, and DOCX formats, including table extraction

import os
import sys
from io import BytesIO
# import PyPDF2   # used pdfplumber instead as it has better text and tabel extraction
import pdfplumber
import docx
from .logger import logging
from .exception import Custom_Exception

class FileHandler:
    @staticmethod
    def read_pdf(file_obj):
        logging.info('Entered read_pdf utils')
        """ 
        Reads the content of a PDF file, including tables.
        
        Parameters:
        - file_obj (BytesIO): File-like object.
        
        Returns:
        - str: Extracted text from the PDF.
        """
        try:
            text=''
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    text+= page.extract_text() or ""
                    tables=page.extract_tables()
                    for table in tables:
                        for row in table:
                            text+= "\t".join(row)+"\n"
            return text if text else "Error:Could not retrieve text"
        except Exception as e:
            raise Custom_Exception(e,sys)
    


    @staticmethod
    def read_docx(file_obj):
        logging.info('Entered read_docx utils')
        """
        Reads the content of a DOCX file, including tables.
        
        Parameters:
        - file_obj (BytesIO): File-like object.
        
        Returns:
        - str: Extracted text from the DOCX file.
        """
        try:
            doc=docx.Document(file_obj)
            text="\n".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    text+= "\t".join(cell.text for cell in row) +"\n"
            return text if text else "Error: Couldn't retrieve text from DOCX"
        except Exception as e:
            raise Custom_Exception(e,sys)

    @staticmethod
    def read_file(file_obj,file_type):
        logging.info('Entered read_file utils')
        """
        Reads the content of a text, PDF, or DOCX file, including tables.
        
        Parameters:
        - file_obj (BytesIO): File-like object.
        - file_type (str): File extension (txt, pdf, docx).
        
        Returns:
        - str: Content of the file.
        """
        try:
            if file_type =='txt':
                return file_obj.read().decode('utf-8')
            elif file_type == 'pdf':
                return FileHandler.read_pdf(file_obj)
            elif file_type == 'docx':
                return FileHandler.read_docx(file_obj)
            else:
                return "Error: Unsupported file type"
        except Exception as e:
            raise Custom_Exception(e,sys)


        except Exception as e:
            raise Custom_Exception(e,sys)